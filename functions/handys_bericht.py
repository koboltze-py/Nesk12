"""
Handys – Schadens-/Verlustbericht
Erstellt ein Word-Dokument für defekte oder verlorene Diensthandys.
"""
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import HANDYS_EXPORT_PATH

_MONATE_DE = {
    1: "01_Januar", 2: "02_Februar", 3: "03_März", 4: "04_April",
    5: "05_Mai",    6: "06_Juni",    7: "07_Juli", 8: "08_August",
    9: "09_September", 10: "10_Oktober", 11: "11_November", 12: "12_Dezember",
}

def berichte_basis_pfad() -> str:
    """Gibt den Basis-Ordner für alle Berichte zurück: Daten/Handys/Berichte"""
    return os.path.join(HANDYS_EXPORT_PATH, "Berichte")

def berichte_ordner_fuer_datum(dt: datetime | None = None) -> str:
    """Gibt den Jahr/Monat-Ordner zurück und erstellt ihn falls nötig."""
    if dt is None:
        dt = datetime.now()
    ordner = os.path.join(
        berichte_basis_pfad(),
        str(dt.year),
        _MONATE_DE[dt.month],
    )
    os.makedirs(ordner, exist_ok=True)
    return ordner

def lade_berichte_fuer_handy(inventarnummer: str) -> list[dict]:
    """
    Durchsucht den Berichte-Ordner und gibt alle Dateien zurück,
    die zur angegebenen Inventarnummer gehören.
    Rückgabe: list of {dateiname, pfad, datum_str, typ}
    """
    basis = berichte_basis_pfad()
    ergebnisse = []
    if not os.path.isdir(basis):
        return ergebnisse
    safe_inv = inventarnummer.replace("/", "-").replace("\\", "-")
    for jahr_dir in sorted(os.listdir(basis), reverse=True):
        jahr_pfad = os.path.join(basis, jahr_dir)
        if not os.path.isdir(jahr_pfad):
            continue
        for monat_dir in sorted(os.listdir(jahr_pfad), reverse=True):
            monat_pfad = os.path.join(jahr_pfad, monat_dir)
            if not os.path.isdir(monat_pfad):
                continue
            for fname in sorted(os.listdir(monat_pfad), reverse=True):
                if not fname.lower().endswith(".docx"):
                    continue
                if safe_inv not in fname:
                    continue
                pfad = os.path.join(monat_pfad, fname)
                # Datum aus Dateiname extrahieren: *_YYYYMMDD_HHMMSS.docx
                datum_str = ""
                try:
                    teile = fname.replace(".docx", "").split("_")
                    # suche 8-stelligen Block
                    for i, t in enumerate(teile):
                        if len(t) == 8 and t.isdigit():
                            dt = datetime.strptime(t, "%Y%m%d")
                            datum_str = dt.strftime("%d.%m.%Y")
                            if i + 1 < len(teile) and len(teile[i+1]) == 6 and teile[i+1].isdigit():
                                zt = teile[i+1]
                                datum_str += f" {zt[:2]}:{zt[2:4]}:{zt[4:]}"
                            break
                except Exception:
                    pass
                typ = "Verlustbericht" if "Verlust" in fname else "Schadensbericht"
                ergebnisse.append({
                    "dateiname": fname,
                    "pfad": pfad,
                    "datum_str": datum_str,
                    "typ": typ,
                    "jahr": jahr_dir,
                    "monat": monat_dir,
                })
    return ergebnisse


def erstelle_schadensbericht(handy: dict, bericht_daten: dict) -> str:
    """
    Erstellt einen Schadens- oder Verlustbericht für ein Diensthandy als Word-Dokument.
    Gibt den Pfad der erstellten Datei zurück.

    handy:         dict mit Feldern aus der handys-Tabelle
    bericht_daten: dict mit Feldern aus dem Dialog:
        ersteller        (str)
        beschreibung     (str)   – überschreibt handy['defekt_beschreibung']
        defekt_datum     (str)   – überschreibt handy['defekt_datum']
        defekt_von       (str)   – überschreibt handy['defekt_gemeldet_von']
        massnahmen       (list[str]) – Liste der angehakten Maßnahmen-Texte
        sonstiges        (str)   – Freitext für Sonstiges
        notizen          (str)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen."
        )

    zustand = handy.get("zustand", "")
    ist_verloren = zustand == "Verloren"
    bericht_typ = "Verlustbericht" if ist_verloren else "Schadensbericht"

    ersteller    = bericht_daten.get("ersteller", "")
    beschreibung = bericht_daten.get("beschreibung", handy.get("defekt_beschreibung", ""))
    defekt_datum = bericht_daten.get("defekt_datum", handy.get("defekt_datum", ""))
    defekt_von   = bericht_daten.get("defekt_von",   handy.get("defekt_gemeldet_von", ""))
    massnahmen   = bericht_daten.get("massnahmen",   [])
    sonstiges    = bericht_daten.get("sonstiges",    "")
    notizen      = bericht_daten.get("notizen",      handy.get("notizen", ""))

    jetzt = datetime.now()
    inv_nr = handy.get("inventarnummer", "unbekannt")
    safe_inv = inv_nr.replace("/", "-").replace("\\", "-")
    stamp = jetzt.strftime("%Y%m%d_%H%M%S")
    dateiname = f"{bericht_typ}_{safe_inv}_{stamp}.docx"

    ordner = berichte_ordner_fuer_datum(jetzt)
    ziel_pfad = os.path.join(ordner, dateiname)

    doc = Document()

    # ── Seitenränder ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    DRK_ROT  = RGBColor(0xCC, 0x00, 0x00)
    DRK_BLAU = RGBColor(0x15, 0x65, 0xA8)
    GRAU     = RGBColor(0x40, 0x40, 0x40)

    def _titel(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DRK_ROT

    def _untertitel(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAU

    def _abschnitt(text: str):
        p = doc.add_paragraph()
        run = p.add_run(f"▌ {text}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DRK_BLAU

    def _trennlinie():
        p = doc.add_paragraph("─" * 72)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    def _zeile(label: str, wert: str):
        p = doc.add_paragraph()
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(wert or "–")
        r_val.font.size = Pt(11)

    def _freitext(text: str):
        for zeile in (text or "").split("\n"):
            p = doc.add_paragraph(zeile or " ")
            if p.runs:
                p.runs[0].font.size = Pt(11)

    def _unterschrift_block():
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            "___________________________          ___________________________\n"
            "Unterschrift Schichtleiter              Unterschrift Geräteverantwortlicher"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = GRAU

    # ── Kopfbereich ───────────────────────────────────────────────────────────
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_p.add_run(
        "Deutsches Rotes Kreuz – Kreisverband Köln e.V.\n"
        "Erste-Hilfe-Station Flughafen Köln/Bonn"
    )
    org_run.font.size = Pt(10)
    org_run.font.color.rgb = GRAU
    doc.add_paragraph()

    _titel(bericht_typ.upper())
    _untertitel("Diensthandy – DRK EHS Flughafen Köln/Bonn")
    _trennlinie()
    doc.add_paragraph()

    # ── 1. Gerätedaten ────────────────────────────────────────────────────────
    _abschnitt("1. Gerätedaten")
    _zeile("Inventarnummer",      handy.get("inventarnummer", ""))
    _zeile("Hersteller",          handy.get("hersteller", ""))
    _zeile("Modell",              handy.get("modell", ""))
    _zeile("Rufnummer",           handy.get("rufnummer", ""))
    _zeile("SIM-Kartennummer",    handy.get("kartennummer", ""))
    _zeile("Standort / Ausgabe",  handy.get("standort", ""))
    _zeile("Anschaffungsdatum",   handy.get("anschaffungsdatum", ""))
    doc.add_paragraph()

    # ── 2. Meldung ────────────────────────────────────────────────────────────
    _abschnitt(f"2. Meldung: {zustand}")
    _zeile("Zustand",             zustand)
    _zeile("Festgestellt am",     defekt_datum)
    _zeile("Festgestellt von",    defekt_von)
    doc.add_paragraph()

    # ── 3. Beschreibung ───────────────────────────────────────────────────────
    _abschnitt("3. Verlustbeschreibung" if ist_verloren else "3. Schadensbeschreibung")
    if beschreibung.strip():
        _freitext(beschreibung.strip())
    else:
        p = doc.add_paragraph("(keine Beschreibung hinterlegt)")
        if p.runs:
            p.runs[0].font.size = Pt(11)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

    # ── 4. Maßnahmen ──────────────────────────────────────────────────────────
    _abschnitt("4. Eingeleitete Maßnahmen")
    if ist_verloren:
        alle_massnahmen = [
            "SIM-Karte beim Anbieter gesperrt",
            "Gerät als verloren gemeldet (Polizei / Fundbüro)",
            "Vorgesetzter informiert",
            "Ersatzgerät ausgegeben",
        ]
    else:
        alle_massnahmen = [
            "Gerät aus dem Dienst genommen",
            "Reparatur veranlasst",
            "Ersatzgerät ausgegeben",
            "Entsorgung / Austausch veranlasst",
        ]
    for m in alle_massnahmen:
        haken = "☑" if m in massnahmen else "☐"
        p = doc.add_paragraph(f"{haken}  {m}")
        if p.runs:
            p.runs[0].font.size = Pt(11)
    # Sonstiges
    sonstiges_haken = "☑" if sonstiges.strip() else "☐"
    sonstiges_text = f"Sonstiges: {sonstiges.strip()}" if sonstiges.strip() else "Sonstiges: _______________________________________"
    p = doc.add_paragraph(f"{sonstiges_haken}  {sonstiges_text}")
    if p.runs:
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ── 5. Notizen ────────────────────────────────────────────────────────────
    _abschnitt("5. Notizen / Weitere Anmerkungen")
    if notizen.strip():
        _freitext(notizen.strip())
    else:
        for _ in range(3):
            p = doc.add_paragraph()
            run = p.add_run("_" * 80)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    doc.add_paragraph()

    # ── Fußbereich ────────────────────────────────────────────────────────────
    _trennlinie()
    _zeile("Bericht erstellt am", jetzt.strftime("%d.%m.%Y %H:%M") + " Uhr")
    if ersteller:
        _zeile("Erstellt von", ersteller)
    doc.add_paragraph()
    _unterschrift_block()

    doc.save(ziel_pfad)
    return ziel_pfad
