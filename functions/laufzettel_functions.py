"""
Laufzettel-Funktionen
Erstellt Word-Dokumente für Mitarbeiter-Laufzettel (Onboarding, Dienstkleidung etc.)
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

from config import BASE_DIR as _BASE_DIR

BASE_DIR = Path(_BASE_DIR)
LAUFZETTEL_DIR = BASE_DIR / "Daten" / "Laufzettel"


# ─── Schulungs-Konfiguration ──────────────────────────────────────────────────

SCHULUNGEN_AUSWAHL = [
    {"key": "vorfeldschulung",  "anzeige": "Vorfeldschulung"},
    {"key": "prm_schulung",     "anzeige": "PRM-Schulung"},
    {"key": "si_schulung",      "anzeige": "SI-Schulung (Sicherheitsschulung)"},
    {"key": "arbeitsschutz",    "anzeige": "Arbeitsschutz"},
    {"key": "erste_hilfe",      "anzeige": "Erste Hilfe"},
]

EINARBEITUNGSTYPEN = [
    "Einarbeitung",
    "Refresher-Schulung",
]


# ─── Hilfsfunktionen ──────────────────────────────────────────────────────────

def _add_heading(doc, text: str, level: int = 1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xa8)
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_checkbox_line(doc, text: str, checked: bool = False):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(16)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    box = "☑" if checked else "☐"
    run = p.add_run(f"{box}  {text}")
    run.font.size = Pt(11)
    return p


def _add_info_line(doc, label: str, value: str = ""):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_val = p.add_run(value if value else "_" * 40)
    run_val.font.size = Pt(11)
    return p


def _add_separator(doc):
    from docx.shared import Pt
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(8)
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _add_unterschrift_zeile(doc, label: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{'_' * 50}   ({label})")
    run.font.size = Pt(10)


# ─── Vorlage: Neu-Einstellung / ZÜP ──────────────────────────────────────────

def _abschnitt_zuep(doc):
    _add_heading(doc, "1.  ZÜP – Zuverlässigkeitsüberprüfung", level=1)
    p = doc.add_paragraph(
        "Für den Einsatz im Sicherheitsbereich des Flughafens Köln/Bonn ist eine "
        "Zuverlässigkeitsüberprüfung (ZÜP) gemäß §7 LuftSiG erforderlich. "
        "Folgende Unterlagen sind einzureichen:"
    )
    from docx.shared import Pt
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)

    unterlagen = [
        "Ausweiskopie (Personalausweis oder Reisepass, beidseitig)",
        "Wohnortmeldebescheinigungen / Nachweise der Wohnorte der letzten 10 Jahre",
        "Nachweise aller Arbeitgeber der letzten 5 Jahre (Arbeitszeugnisse, Arbeitsverträge)",
        "Letzte Lohnabrechnung (als Tätigkeitsnachweis)",
        "Sonstige Nachweise (z. B. Ausbildungsnachweise, Zeugnisse)",
    ]
    for u in unterlagen:
        _add_checkbox_line(doc, u)

    doc.add_paragraph()
    _add_info_line(doc, "Ansprechpartner", "Schichtleiter sowie Herr Peters")
    _add_info_line(doc, "Eingereicht am", "")
    _add_info_line(doc, "Geprüft von", "")
    _add_info_line(doc, "Bemerkungen", "")


# ─── Vorlage: Dienstkleidung ──────────────────────────────────────────────────

def _abschnitt_dienstkleidung(doc):
    _add_heading(doc, "2.  Dienstkleidung", level=1)
    p = doc.add_paragraph(
        "Der/Die Mitarbeiter/in benötigt Dienstkleidung gemäß DRK-Vorgabe. "
        "Bitte folgende Schritte beachten:"
    )
    from docx.shared import Pt
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)

    schritte = [
        "Kleidungsgrößen angeben (Oberteil, Hose, Schuhgröße)",
        "Ausgabe durch Herrn Etz oder Herrn Kurthen (oder vertretungsweise Schichtleiter)",
        "Quittierung der ausgegebenen Kleidungsstücke",
        "Rückgabe beim Ausscheiden vollständig klären",
    ]
    for s in schritte:
        _add_checkbox_line(doc, s)

    doc.add_paragraph()
    _add_info_line(doc, "Hauptansprechpartner", "Herr Etz, Herr Kurthen")
    _add_info_line(doc, "Weitere Ansprechpartner", "Alle Schichtleiter")
    _add_info_line(doc, "Ausgabe erfolgt am", "")
    _add_info_line(doc, "Ausgegeben von", "")
    _add_info_line(doc, "Bemerkungen", "")


# ─── Abschnitt: Schulungen ────────────────────────────────────────────────────

def _abschnitt_schulungen(doc, schulungen: list[str], einarbeitung_typ: str | None, num: int):
    if not schulungen and not einarbeitung_typ:
        return

    _add_heading(doc, f"{num}.  Schulungen", level=1)
    p = doc.add_paragraph(
        "Folgende Schulungen müssen absolviert werden. "
        "Bitte Termine mit den zuständigen Ansprechpartnern vereinbaren:"
    )
    from docx.shared import Pt
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)

    schulungs_details = {
        "Vorfeldschulung":                ("Vorfeldschulung (Pflicht für Vorfeldbewegungen)",  "Koordination über Schichtleiter / extern"),
        "PRM-Schulung":                   ("PRM-Schulung (Persons with Reduced Mobility)",      "Koordination über Schichtleiter"),
        "SI-Schulung (Sicherheitsschulung)": ("SI-Schulung / Sicherheitsschulung",              "Flughafensicherheit – Terminvereinbarung erforderlich"),
        "Arbeitsschutz":                  ("Arbeitsschutz-Unterweisung",                        "Koordination über Schichtleiter / intern"),
        "Erste Hilfe":                    ("Erste-Hilfe-Kurs (16h)",                            "Intern – Terminvereinbarung über Schichtleiter"),
    }

    for sch in schulungen:
        label, info = schulungs_details.get(sch, (sch, ""))
        _add_checkbox_line(doc, label)
        if info:
            p2 = doc.add_paragraph(f"     → {info}")
            p2.paragraph_format.left_indent = Pt(30)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(4)
            for run in p2.runs:
                run.font.size = Pt(10)
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _add_info_line(doc, "   Termin", "")
        _add_info_line(doc, "   Absolviert am", "")

    if einarbeitung_typ:
        _add_checkbox_line(doc, einarbeitung_typ)
        _add_info_line(doc, "   Termin", "")
        _add_info_line(doc, "   Absolviert am", "")
        _add_info_line(doc, "   Eingearbeitet durch", "")


# ─── Hauptfunktion ────────────────────────────────────────────────────────────

def erstelle_laufzettel_word(daten: dict) -> str:
    """
    Erstellt ein Word-Dokument für einen Mitarbeiter-Laufzettel.

    Erwartete Schlüssel in ``daten``:
        mitarbeiter_name   – vollständiger Name des Mitarbeiters
        datum              – Datum (str im Format DD.MM.YYYY), optional
        vorlagen           – list[str]: z. B. ["Neu-Einstellung", "Dienstkleidung"]
        schulungen         – list[str]: ausgewählte Schulungsnamen
        einarbeitung_typ   – str | None: "Einarbeitung" oder "Refresher-Schulung"
        erstellt_von       – Name des Erstellers (optional)

    Gibt den vollständigen Dateipfad zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx ist nicht installiert. "
            "Bitte 'pip install python-docx' ausführen."
        )

    LAUFZETTEL_DIR.mkdir(parents=True, exist_ok=True)

    ma_name     = daten.get("mitarbeiter_name", "Unbekannt")
    datum       = daten.get("datum", datetime.now().strftime("%d.%m.%Y"))
    vorlagen    = daten.get("vorlagen", [])
    schulungen  = daten.get("schulungen", [])
    einarbeitung_typ = daten.get("einarbeitung_typ", None)
    erstellt_von = daten.get("erstellt_von", "")

    doc = Document()

    # ── Seitenränder ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Kopfzeile ─────────────────────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_para.add_run("LAUFZETTEL – MITARBEITER ONBOARDING")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xa8)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_para.add_run("DRK – Erste-Hilfe-Station Flughafen Köln/Bonn")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    _add_separator(doc)

    # ── Mitarbeiterdaten ──────────────────────────────────────────────────────
    _add_heading(doc, "Mitarbeiterdaten", level=1)
    _add_info_line(doc, "Name", ma_name)
    _add_info_line(doc, "Eintrittsdatum", datum)
    _add_info_line(doc, "Erstellt am", datetime.now().strftime("%d.%m.%Y"))
    if erstellt_von:
        _add_info_line(doc, "Erstellt von", erstellt_von)

    _add_separator(doc)

    # ── Vorlagen-Abschnitte ───────────────────────────────────────────────────
    num = 1
    if "Neu-Einstellung" in vorlagen:
        _abschnitt_zuep(doc)
        _add_separator(doc)
        num += 1

    if "Dienstkleidung" in vorlagen:
        # Nummerierung anpassen
        _abschnitt_dienstkleidung(doc)
        _add_separator(doc)
        num += 1

    # ── Schulungen-Abschnitt ──────────────────────────────────────────────────
    _abschnitt_schulungen(doc, schulungen, einarbeitung_typ, num)

    if schulungen or einarbeitung_typ:
        _add_separator(doc)

    # ── Unterschriften ────────────────────────────────────────────────────────
    _add_heading(doc, "Bestätigungen", level=1)
    p_info = doc.add_paragraph(
        "Der/Die Mitarbeiter/in bestätigt mit Unterschrift die Kenntnisnahme "
        "und die ordnungsgemäße Erledigung der aufgeführten Punkte."
    )
    for run in p_info.runs:
        run.font.size = Pt(11)
    p_info.paragraph_format.space_after = Pt(8)

    _add_unterschrift_zeile(doc, "Mitarbeiter/in")
    _add_unterschrift_zeile(doc, "Schichtleiter/in")
    _add_unterschrift_zeile(doc, "Datum")

    # ── Datei speichern ───────────────────────────────────────────────────────
    safe_name = "".join(
        c for c in ma_name if c.isalnum() or c in " _-"
    ).strip().replace(" ", "_")[:40]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dateiname = f"Laufzettel_{safe_name}_{ts}.docx"
    ziel_pfad = LAUFZETTEL_DIR / dateiname

    doc.save(str(ziel_pfad))
    return str(ziel_pfad)


def oeffne_laufzettel(pfad: str) -> None:
    """Öffnet das Laufzettel-Dokument mit der Standard-Anwendung."""
    import subprocess
    import os
    if os.path.isfile(pfad):
        os.startfile(pfad)
