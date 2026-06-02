"""
Workflow-Widget – Abgleich Stärkemeldungen ↔ Tagesdienstpläne
==============================================================
Lädt beliebig viele Stärkemeldungen (Word .docx) und Tagesdienstpläne
(Excel .xlsx) und vergleicht die enthaltenen Personen / Dienste / Zeiten.
"""
from __future__ import annotations

import os
import re
from difflib import get_close_matches
from pathlib import Path
from datetime import datetime

from PySide6.QtCore import Qt, QThread, QUrl, Signal
from PySide6.QtGui import QDesktopServices, QFont, QColor
from PySide6.QtWidgets import (
    QCheckBox, QComboBox, QDialog, QDialogButtonBox, QFileDialog, QFrame,
    QGroupBox, QHBoxLayout, QHeaderView, QLabel, QMenu, QMessageBox,
    QPushButton, QScrollArea, QSizePolicy, QSpinBox, QSplitter,
    QTableWidget, QTableWidgetItem, QTextEdit, QVBoxLayout, QWidget,
    QProgressBar,
)

from config import BASE_DIR, FIORI_BLUE, FIORI_TEXT
from database.workflow_db import (
    lade_eintrag, speichere_eintrag, init_db,
    alle_monate, lade_session, speichere_session, loesche_session,
)

# ── Standardpfade ──────────────────────────────────────────────────────────────
_BASE_ONEDRIVE = Path(BASE_DIR).parent.parent  # …/!Gemeinsam.26/
_DEFAULT_STAERKE  = str(_BASE_ONEDRIVE / "06_Stärkemeldung")
_DEFAULT_DIENSTPLAN = str(_BASE_ONEDRIVE / "04_Tagesdienstpläne")


# ── Stärkemeldung-Parser (Word .docx) ─────────────────────────────────────────

# Zeitraum-Regex: HH:MM-HH:MM, gefolgt von Tab oder mind. 2 Leerzeichen
_SM_ZEIT_RE = re.compile(r'^(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})\s*(?:\t|\s{2,})(.*)')

# Abschnitt-Header in der Stärkemeldung (lowercase)
_SM_ABSCHNITTE: dict[str, str] = {
    "schichtleiter":      "SL",
    "disposition":        "Dispo",
    "behindertenbetreuer": "Betreuer",
}


def _datum_aus_dateiname(stem: str) -> str:
    """Extrahiert Datum aus Dateinamen; normiert auf YYYY-MM-DD."""
    # Format DD.MM.YYYY (z.B. "01.05.2026")
    m = re.search(r'(\d{1,2})\.(\d{2})\.(\d{4})', stem)
    if m:
        return f"{m.group(3)}-{m.group(2)}-{int(m.group(1)):02d}"
    # Format YYYY-MM-DD oder YYYY_MM_DD
    m2 = re.search(r'(\d{4})[-_](\d{2})[-_](\d{2})', stem)
    if m2:
        return f"{m2.group(1)}-{m2.group(2)}-{m2.group(3)}"
    return ""


def _parse_staerkemeldung(docx_path: str) -> dict:
    """
    Liest Personen aus einer Stärkemeldungs-.docx im DRK-Standardformat.
    Format: 1 äußere Tabelle, 2 Spalten; Daten als Absätze in der rechten Zelle.
    Abschnitte: Schichtleiter / Disposition / Behindertenbetreuer
    Zeilen: HH:MM-HH:MM<Tab>Nachname1 / Nachname2 ...

    Rückgabe: { 'datei', 'datum', 'personen': list[dict] }
    Person: { 'vollname', 'nachname', 'dienst' (SL|Dispo|Betreuer), 'beginn', 'ende' }
    """
    try:
        from docx import Document
    except ImportError:
        return {"datei": Path(docx_path).name, "datum": "", "personen": [], "fehler": "python-docx fehlt"}

    try:
        doc = Document(docx_path)
    except Exception as e:
        return {"datei": Path(docx_path).name, "datum": "", "personen": [], "fehler": str(e)}

    datum_str = _datum_aus_dateiname(Path(docx_path).stem)
    personen: list[dict] = []

    # Rechte Zelle der äußeren Tabelle enthält alle Daten
    if not doc.tables:
        return {"datei": Path(docx_path).name, "datum": datum_str, "personen": personen}

    right_cell = doc.tables[0].rows[0].cells[-1]
    paragraphen = right_cell.paragraphs

    aktueller_abschnitt: str | None = None

    for para in paragraphen:
        text = para.text.strip()
        if not text:
            continue

        # Abschnitt-Header?
        text_lower = text.lower()
        for key, val in _SM_ABSCHNITTE.items():
            if text_lower.startswith(key):
                aktueller_abschnitt = val
                break
        else:
            # Kein Header-Match – Datenzeile?
            if aktueller_abschnitt is None:
                continue
            m = _SM_ZEIT_RE.match(text)
            if not m:
                continue

            beginn   = m.group(1).strip()
            ende     = m.group(2).strip()
            namen_raw = m.group(3).strip()

            # Mehrere Namen: " / " als Trennzeichen
            namen = [n.strip() for n in namen_raw.split(" / ") if n.strip()]

            for name_raw in namen:
                # Nachname-Extraktion aus SM-Format:
                # - "Groß"           → nachname="groß"
                # - "Blei Pa"        → letztes Token ≤2 Zeichen = Vorname-Abk. → nachname="blei"
                # - "Tepealan Le"    → nachname="tepealan"
                # - "El Mojahid"     → letztes Token >2 Zeichen = Compound → nachname="el mojahid"
                # - "Clausen-Hansen Be" → nachname="clausen-hansen"
                tokens = name_raw.split()
                if len(tokens) == 1:
                    nachname = tokens[0].lower()
                    abbrev   = ""
                elif len(tokens[-1]) <= 2:
                    # Letztes Token ist Vorname-Abkürzung (z.B. "Blei Pa" → abbrev="pa")
                    nachname = " ".join(tokens[:-1]).lower()
                    abbrev   = tokens[-1].lower()
                else:
                    # Compound-Nachname (z.B. "El Mojahid")
                    nachname = name_raw.lower()
                    abbrev   = ""
                personen.append({
                    "vollname": name_raw,
                    "nachname": nachname,
                    "abbrev":   abbrev,
                    "dienst":   aktueller_abschnitt,
                    "beginn":   beginn,
                    "ende":     ende,
                })

    # Fallback: Datum aus Dokumenttext
    if not datum_str:
        for para in doc.paragraphs:
            m = re.search(r'(\d{1,2})\.(\d{2})\.(\d{4})', para.text)
            if m:
                datum_str = f"{m.group(3)}-{m.group(2)}-{int(m.group(1)):02d}"
                break

    return {
        "datei":    Path(docx_path).name,
        "datum":    datum_str,
        "personen": personen,
    }


# ── Dienstplan-Parser (Excel .xlsx) ───────────────────────────────────────────

def _parse_dienstplan_fuer_abgleich(xlsx_path: str) -> dict:
    """
    Nutzt den vorhandenen DienstplanParser und gibt eine vereinfachte
    Personen-Liste zurück.
    """
    try:
        from functions.dienstplan_parser import DienstplanParser
        result = DienstplanParser(xlsx_path, alle_anzeigen=True, round_dispo=False).parse()
    except Exception as e:
        return {"datei": xlsx_path, "datum": "", "personen": [], "fehler": str(e)}

    if not result.get("success", True) and result.get("error"):
        return {
            "datei":    Path(xlsx_path).name,
            "datum":    "",
            "personen": [],
            "fehler":   result["error"],
        }

    personen: list[dict] = []
    for p in result.get("betreuer", []) + result.get("dispo", []):
        # Korrekter Key aus DienstplanParser ist "dienst_kategorie" (z.B. T, N, DT, DN)
        dienst   = p.get("dienst_kategorie") or ""
        # R = Rufbereitschaft → wird nicht in der SM erfasst, daher beim Abgleich ignorieren
        if dienst.strip().upper() == "R":
            continue
        # Lars Peters wird generell nicht verglichen
        vollname_lower = (p.get("vollname") or "").strip().lower()
        if vollname_lower == "lars peters":
            continue
        nachname = (p.get("nachname") or "").strip().lower()
        personen.append({
            "vollname": p.get("vollname", "").strip(),
            "nachname": nachname,
            "dienst":   dienst.strip(),
            "beginn":   (p.get("start_zeit") or "").strip(),
            "ende":     (p.get("end_zeit")   or "").strip(),
            "ist_dispo": bool(p.get("ist_dispo")),
            "notiz":    (p.get("notiz") or "").strip(),
        })

    # Datum aus Dateiname (DD.MM.YYYY oder YYYY-MM-DD)
    datum_str = _datum_aus_dateiname(Path(xlsx_path).stem)

    return {
        "datei":    Path(xlsx_path).name,
        "datum":    datum_str,
        "personen": personen,
    }


# ── Abgleich-Logik ─────────────────────────────────────────────────────────────

def _normiere_name(s: str) -> str:
    return " ".join(s.lower().split())

def _normiere_zeit(s: str) -> str:
    """Normiert Zeiten auf HH:MM (volle Stunde wenn nur Stunde angegeben)."""
    s = s.strip().replace(".", ":")
    if re.fullmatch(r'\d{1,2}', s):
        return f"{int(s):02d}:00"
    m = re.match(r'(\d{1,2}):(\d{2})', s)
    if m:
        return f"{int(m.group(1)):02d}:{m.group(2)}"
    return s

def _normiere_dienst(s: str) -> str:
    return s.strip().upper()


def _stunde_von(zeit: str) -> str:
    """Gibt nur die Stunden-Komponente als 'HH:00' zurück, z.B. '07:30' → '07:00'."""
    m = re.match(r'(\d{1,2}):\d{2}', zeit)
    return f"{int(m.group(1)):02d}:00" if m else zeit


def _zeiten_stimmen_ueberein(sm_zeit: str, dp_zeit: str, ist_sl_dispo: bool) -> bool:
    """
    Vergleicht SM- und DP-Zeit.
    Für SL/Dispo: SM-Zeiten sind auf die volle Stunde nach unten gerundet.
    → Übereinstimmung wenn SM-Zeit == floor(DP-Zeit).
    Für Betreuer: exakter Vergleich.
    """
    if not sm_zeit or not dp_zeit:
        return True  # fehlende Zeit nicht als Abweichung werten
    if sm_zeit == dp_zeit:
        return True
    if ist_sl_dispo:
        return _stunde_von(dp_zeit) == sm_zeit
    return False


class AbgleichErgebnis:
    """Enthält alle Unterschiede zwischen Stärkemeldung und Dienstplan."""

    def __init__(self, datei_staerke: str, datei_dienstplan: str,
                 pfad_staerke: str = "", pfad_dienstplan: str = "", datum: str = ""):
        self.datei_staerke    = datei_staerke
        self.datei_dienstplan = datei_dienstplan
        self.pfad_staerke     = pfad_staerke
        self.pfad_dienstplan  = pfad_dienstplan
        self.datum            = datum        # YYYY-MM-DD des SM-Startdatums
        self.ok:         list[dict] = []   # identisch
        self.abweichung: list[dict] = []   # Unterschied im Dienst/Zeit
        self.nur_staerke: list[dict] = []  # in Stärke, nicht im Dienstplan
        self.nur_dienstplan: list[dict] = []  # im Dienstplan, nicht in Stärke

    @property
    def hat_fehler(self) -> bool:
        return bool(self.abweichung or self.nur_staerke or self.nur_dienstplan)


def _sm_kat_zu_norm(sm_dienst: str) -> str:
    """Normiert SM-Abschnitt auf 'Dispo' oder 'Betreuer' für den Vergleich."""
    return "Dispo" if sm_dienst in ("SL", "Dispo") else "Betreuer"


def _dp_kat_zu_norm(ist_dispo: bool) -> str:
    return "Dispo" if ist_dispo else "Betreuer"


def _abgleichen(
    staerke_data: dict,
    dienstplan_data: dict,
    pfad_staerke: str = "",
    pfad_dienstplan: str = "",
) -> AbgleichErgebnis:
    """
    Gleicht SM-Personen mit DP-Personen ab.
    Matching erfolgt über den Nachnamen (SM hat nur Nachname).
    Bei Namensgleichheit werden Zeiten und Kategorie (Dispo/Betreuer) verglichen.
    Fuzzy-Fallback (difflib, cutoff 0.82) für Tippfehler in DP-Namen.
    """
    erg = AbgleichErgebnis(
        staerke_data.get("datei", ""),
        dienstplan_data.get("datei", ""),
        pfad_staerke=pfad_staerke,
        pfad_dienstplan=pfad_dienstplan,
        datum=staerke_data.get("datum", ""),
    )

    # Index aufbauen: nachname (lower) → list[dict]
    s_by_nn: dict[str, list[dict]] = {}
    for p in staerke_data.get("personen", []):
        nn = p.get("nachname") or ""
        if not nn and p.get("vollname"):
            nn = p["vollname"].split()[0].lower()
        if nn:
            s_by_nn.setdefault(nn, []).append(p)

    d_by_nn: dict[str, list[dict]] = {}
    for p in dienstplan_data.get("personen", []):
        nn = (p.get("nachname") or "").lower().strip()
        if not nn:
            parts = p.get("vollname", "").split()
            nn = parts[-1].lower() if parts else ""
        if nn:
            d_by_nn.setdefault(nn, []).append(p)

    # Pre-resolve 1a: SM compound-Nachnamen ohne DP-Match → letztes Token versuchen.
    # DP-Parser speichert manchmal nur das letzte Wort als Nachname ("Mojahid"),
    # SM schreibt den vollen Compound-Namen ("El Mojahid").
    for nn_c in list(s_by_nn.keys()):
        if " " in nn_c and nn_c not in d_by_nn:
            last = nn_c.split()[-1]
            if last in d_by_nn and last not in s_by_nn:
                for p in s_by_nn[nn_c]:
                    s_by_nn.setdefault(last, []).append(p)
                del s_by_nn[nn_c]

    # Pre-resolve 1b: DP compound-Nachnamen → SM single-Token-Match.
    # DP hat "el mojahid" als nachname, SM schreibt nur "mojahid".
    # → SM-Eintrag unter "mojahid" auf den DP-Key "el mojahid" remappen.
    # Gleiches gilt für erstes Token: DP "moeeni mahvelati", SM "moeeni".
    for nn_d in list(d_by_nn.keys()):
        if " " in nn_d and nn_d not in s_by_nn:
            parts_d = nn_d.split()
            for token in (parts_d[-1], parts_d[0]):  # letztes UND erstes Token prüfen
                if token in s_by_nn and token not in d_by_nn:
                    for p in s_by_nn[token]:
                        s_by_nn.setdefault(nn_d, []).append(p)
                    del s_by_nn[token]
                    break  # nur einmal remappen

    # Pre-resolve 2: Fuzzy-Matching für SM-Nachnamen ohne exakten DP-Treffer
    # (Tippfehler im Dienstplan, z.B. "Müler" statt "Müller").
    dp_nn_keys = list(d_by_nn.keys())
    for nn_s in list(s_by_nn.keys()):
        if nn_s not in d_by_nn:
            matches = get_close_matches(nn_s, dp_nn_keys, n=1, cutoff=0.82)
            if matches:
                fuzzy_nn = matches[0]
                # Nur remappen wenn DP-Key noch nicht direkt von SM belegt
                if fuzzy_nn not in s_by_nn:
                    for p in s_by_nn[nn_s]:
                        # Tippfehler-Hinweis im vollname ergänzen
                        p = dict(p)
                        dp_sample = d_by_nn[fuzzy_nn][0]
                        p["vollname"] = p.get("vollname", nn_s) + f" [≈ {dp_sample.get('vollname', fuzzy_nn)}]"
                        s_by_nn.setdefault(fuzzy_nn, []).append(p)
                    del s_by_nn[nn_s]

    alle_nn = set(s_by_nn) | set(d_by_nn)

    for nn in sorted(alle_nn):
        sp_liste = s_by_nn.get(nn, [])
        dp_liste = d_by_nn.get(nn, [])

        if sp_liste and dp_liste:
            # Paare bilden (1:1 in Reihenfolge, bei Mehrfach-Nachnamen)
            # Einträge MIT Vorname-Abkürzung zuerst: sie sind spezifischer und
            # sollen nicht durch einen Fallback-Match verdrängt werden.
            sp_liste_sorted = sorted(sp_liste, key=lambda x: (0 if x.get("abbrev") else 1))
            gematchte_dp: set[int] = set()

            for sp in sp_liste_sorted:
                abbrev = (sp.get("abbrev") or "").lower()

                # Besten DP-Match suchen: erst per Vorname-Abkürzung, dann first-unmatched
                dp = None
                if abbrev:
                    for d in dp_liste:
                        if id(d) in gematchte_dp:
                            continue
                        # DP-Vollname ist "Vorname Nachname" → erstes Token ist Vorname
                        vn_parts = d.get("vollname", "").split()
                        vorname = vn_parts[0].lower() if vn_parts else ""
                        if vorname.startswith(abbrev):
                            dp = d
                            break
                if dp is None:
                    # Fallback: erster noch-nicht-gematchter Eintrag
                    dp = next((d for d in dp_liste if id(d) not in gematchte_dp), None)
                if dp is None:
                    erg.nur_staerke.append({
                        "name":      sp.get("vollname", nn),
                        "sm_dienst": sp.get("dienst", ""),
                        "sm_beginn": sp.get("beginn", ""),
                        "sm_ende":   sp.get("ende",   ""),
                    })
                    continue
                gematchte_dp.add(id(dp))

                s_beginn = _normiere_zeit(sp.get("beginn", ""))
                d_beginn = _normiere_zeit(dp.get("beginn", ""))
                s_ende   = _normiere_zeit(sp.get("ende",   ""))
                d_ende   = _normiere_zeit(dp.get("ende",   ""))

                sm_kat = _sm_kat_zu_norm(sp.get("dienst", ""))
                dp_kat = _dp_kat_zu_norm(dp.get("ist_dispo", False))
                ist_sl_dispo = sm_kat == "Dispo"  # SL und Dispo werden gerundet

                unterschiede: list[str] = []
                if sm_kat != dp_kat:
                    unterschiede.append(f"Kategorie: SM={sm_kat} / DP={dp_kat}")

                # Zeiten vergleichen – SL/Dispo toleriert Rundung auf volle Stunde
                if not _zeiten_stimmen_ueberein(s_beginn, d_beginn, ist_sl_dispo):
                    unterschiede.append(f"Beginn: SM={s_beginn} / DP={d_beginn}")
                if not _zeiten_stimmen_ueberein(s_ende, d_ende, ist_sl_dispo):
                    unterschiede.append(f"Ende: SM={s_ende} / DP={d_ende}")

                eintrag = {
                    "name":      dp.get("vollname") or sp.get("vollname", nn),
                    "sm_dienst": sp.get("dienst",  ""),
                    "sm_beginn": sp.get("beginn",  ""),
                    "sm_ende":   sp.get("ende",    ""),
                    "dp_dienst": dp.get("dienst",  ""),
                    "dp_beginn": dp.get("beginn",  ""),
                    "dp_ende":   dp.get("ende",    ""),
                    "dp_notiz":  dp.get("notiz",   ""),
                    "unterschiede": unterschiede,
                }
                if unterschiede:
                    erg.abweichung.append(eintrag)
                else:
                    erg.ok.append(eintrag)

            # Übrige DP-Einträge ohne SM-Partner
            for dp in dp_liste:
                if id(dp) not in gematchte_dp:
                    erg.nur_dienstplan.append({
                        "name":      dp.get("vollname", nn),
                        "dp_dienst": dp.get("dienst",  ""),
                        "dp_beginn": dp.get("beginn",  ""),
                        "dp_ende":   dp.get("ende",    ""),
                        "dp_notiz":  dp.get("notiz",   ""),
                    })

        elif sp_liste:
            for sp in sp_liste:
                erg.nur_staerke.append({
                    "name":      sp.get("vollname", nn),
                    "sm_dienst": sp.get("dienst", ""),
                    "sm_beginn": sp.get("beginn", ""),
                    "sm_ende":   sp.get("ende",   ""),
                })
        else:
            for dp in dp_liste:
                erg.nur_dienstplan.append({
                    "name":      dp.get("vollname", nn),
                    "dp_dienst": dp.get("dienst",  ""),
                    "dp_beginn": dp.get("beginn",  ""),
                    "dp_ende":   dp.get("ende",    ""),
                    "dp_notiz":  dp.get("notiz",   ""),
                })

    return erg


# ── Lade-Thread ────────────────────────────────────────────────────────────────

class _LadeThread(QThread):
    fortschritt = Signal(str)
    fertig      = Signal(list)   # list[AbgleichErgebnis]
    fehler      = Signal(str)

    def __init__(self, staerke_pfade: list[str], dienstplan_pfade: list[str]):
        super().__init__()
        self._staerke    = staerke_pfade
        self._dienstplan = dienstplan_pfade

    def run(self):
        try:
            ergebnisse: list[AbgleichErgebnis] = []

            # Daten laden
            self.fortschritt.emit("Stärkemeldungen werden geladen …")
            sm_data: list[dict] = []
            for pfad in self._staerke:
                self.fortschritt.emit(f"  Lese: {Path(pfad).name}")
                d = _parse_staerkemeldung(pfad)
                d["_pfad"] = pfad
                sm_data.append(d)

            self.fortschritt.emit("Dienstpläne werden geladen …")
            dp_data: list[dict] = []
            for pfad in self._dienstplan:
                self.fortschritt.emit(f"  Lese: {Path(pfad).name}")
                d2 = _parse_dienstplan_fuer_abgleich(pfad)
                d2["_pfad"] = pfad
                dp_data.append(d2)

            # Abgleich: versuche Datum-Matching
            self.fortschritt.emit("Abgleich wird durchgeführt …")

            # Datum-Index der Dienstpläne aufbauen
            dp_by_datum: dict[str, list[dict]] = {}
            for dp in dp_data:
                d = dp.get("datum", "")
                dp_by_datum.setdefault(d, []).append(dp)

            matched_dp: set[int] = set()

            for sm in sm_data:
                sm_datum = sm.get("datum", "")
                # Strikt datum-basiertes Matching: Start-Datum der SM bestimmt den DP
                candidates = dp_by_datum.get(sm_datum, []) if sm_datum else []
                unmatch = [c for c in candidates if id(c) not in matched_dp]
                if unmatch:
                    dp_match = unmatch[0]
                    matched_dp.add(id(dp_match))
                elif not sm_datum:
                    # Kein Datum extrahierbar – mit leerem Platzhalter anzeigen
                    dp_match = {"datei": "(kein Datum)", "datum": "", "personen": []}
                else:
                    # Datum vorhanden, aber kein passender DP gefunden
                    dp_match = {"datei": f"(kein DP f\u00fcr {sm_datum})", "datum": sm_datum, "personen": []}

                ergebnisse.append(_abgleichen(
                    sm, dp_match,
                    pfad_staerke=sm.get("_pfad", ""),
                    pfad_dienstplan=dp_match.get("_pfad", ""),
                ))

            # Übriggebliebene Dienstpläne ohne Stärkemeldung
            for dp in dp_data:
                if id(dp) not in matched_dp:
                    sm_leer = {"datei": "–", "datum": "", "personen": []}
                    ergebnisse.append(_abgleichen(
                        sm_leer, dp,
                        pfad_staerke="",
                        pfad_dienstplan=dp.get("_pfad", ""),
                    ))

            self.fertig.emit(ergebnisse)
        except Exception as exc:
            self.fehler.emit(str(exc))


# ── Detail-Dialog ──────────────────────────────────────────────────────────────

class _DetailDialog(QDialog):
    def __init__(self, erg: AbgleichErgebnis, parent=None):
        super().__init__(parent)
        self._erg = erg
        self.setWindowTitle(f"🔍  Detail: {erg.datei_staerke} ↔ {erg.datei_dienstplan}")
        self.resize(1200, 620)
        lay = QVBoxLayout(self)
        lay.setContentsMargins(12, 10, 12, 10)

        # Statistik-Zeile
        stats = QLabel(
            f"✅ Identisch: {len(erg.ok)}   "
            f"⚠️ Abweichungen: {len(erg.abweichung)}   "
            f"📄 Nur Stärke: {len(erg.nur_staerke)}   "
            f"📋 Nur Dienstplan: {len(erg.nur_dienstplan)}"
        )
        stats.setStyleSheet("font-size: 12px; font-weight: bold; color: #2c3e50; padding: 4px 0;")
        lay.addWidget(stats)

        tbl = QTableWidget()
        tbl.setColumnCount(9)
        tbl.setHorizontalHeaderLabels([
            "Status", "Name",
            "SM Dienst", "SM Beginn", "SM Ende",
            "DP Dienst", "DP Beginn", "DP Ende",
            "📋 Notiz (DP)",
        ])
        tbl.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        tbl.horizontalHeader().setSectionResizeMode(8, QHeaderView.ResizeMode.Stretch)
        tbl.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        tbl.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        tbl.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        tbl.setAlternatingRowColors(True)
        tbl.verticalHeader().setVisible(False)
        tbl.setStyleSheet("font-size: 11px;")
        tbl.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        tbl.customContextMenuRequested.connect(
            lambda pos: self._kontext_menu(tbl, pos)
        )

        rows: list[tuple[str, dict, QColor]] = []
        for e in erg.abweichung:
            rows.append(("⚠️ Abweichung", e, QColor("#fff3cd")))
        for e in erg.nur_staerke:
            rows.append(("📄 Nur Stärke", e, QColor("#fde8e8")))
        for e in erg.nur_dienstplan:
            rows.append(("📋 Nur Dienstplan", e, QColor("#e8f4fd")))
        for e in erg.ok:
            rows.append(("✅ Identisch", e, QColor("#f0fff0")))

        def _ist_dispo_sl(eintrag: dict) -> bool:
            """True wenn SM-Dienst SL oder Dispo ist."""
            d = (eintrag.get("sm_dienst") or eintrag.get("dp_dienst") or "").upper()
            return d in ("SL", "DISPO", "DT", "DN", "DT3", "DN3", "DN10", "DT10")

        # Sortiergewichte Betreuer: T, T10, T8, N, N10, Rest
        _BETR_ORDER = {"T": 0, "T10": 1, "T8": 2, "N": 3, "N10": 4}
        # Sortiergewichte Dispo/SL: SL+DT3, SL+DN3, Dispo+DT, DN, Rest
        def _dispo_key(e: dict) -> int:
            sm = (e.get("sm_dienst") or "").upper()
            dp = (e.get("dp_dienst") or "").upper()
            if sm == "SL" and dp == "DT3":  return 0
            if sm == "SL" and dp == "DN3":  return 1
            if dp in ("DT3",):              return 0  # nur DP bekannt
            if dp in ("DN3",):              return 1
            if dp == "DT":                  return 2
            if dp == "DN":                  return 3
            if sm in ("SL",):               return 4
            return 9

        def _betr_key(e: dict) -> int:
            dp = (e.get("dp_dienst") or e.get("sm_dienst") or "").upper()
            return _BETR_ORDER.get(dp, 99)

        # Sortierung: Dispo/SL zuerst (nach Gewicht), dann Betreuer (nach Gewicht)
        rows_dispo = sorted(
            [(s, e, bg) for s, e, bg in rows if _ist_dispo_sl(e)],
            key=lambda r: _dispo_key(r[1])
        )
        rows_betr = sorted(
            [(s, e, bg) for s, e, bg in rows if not _ist_dispo_sl(e)],
            key=lambda r: _betr_key(r[1])
        )

        # Trennzeilen-Sentinel: (None, None, None)
        combined: list = []
        if rows_dispo:
            combined += rows_dispo
        if rows_dispo and rows_betr:
            combined.append(None)   # Trennzeile
        combined += rows_betr

        tbl.setRowCount(len(combined))
        for ri, row in enumerate(combined):
            if row is None:
                # Trennzeile: grauer Balken mit Label
                sep = QTableWidgetItem("── Betreuer ──────────────────────────────────────────")
                sep.setBackground(QColor("#d5d8dc"))
                sep.setForeground(QColor("#5d6d7e"))
                sep.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
                sep.setFlags(Qt.ItemFlag.NoItemFlags)
                tbl.setItem(ri, 0, sep)
                tbl.setSpan(ri, 0, 1, 9)
                tbl.setRowHeight(ri, 18)
                continue

            status, e, bg = row
            vals = [
                status,
                e.get("name", ""),
                e.get("sm_dienst", ""),
                e.get("sm_beginn", ""),
                e.get("sm_ende",   ""),
                e.get("dp_dienst", ""),
                e.get("dp_beginn", ""),
                e.get("dp_ende",   ""),
                e.get("dp_notiz",  ""),
            ]
            for ci, v in enumerate(vals):
                item = QTableWidgetItem(v)
                item.setBackground(bg)
                if ci == 0 and "Abweichung" in status:
                    item.setForeground(QColor("#856404"))
                elif ci == 0 and "Stärke" in status:
                    item.setForeground(QColor("#c0392b"))
                elif ci == 0 and "Dienstplan" in status:
                    item.setForeground(QColor("#1a6ea8"))
                if ci == 2 and _ist_dispo_sl(e):
                    item.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
                if ci == 8 and v:
                    item.setForeground(QColor("#555"))
                    item.setToolTip(v)
                tbl.setItem(ri, ci, item)

            # Unterschiede-Tooltip
            if e.get("unterschiede"):
                for ci in range(8):
                    if tbl.item(ri, ci):
                        tbl.item(ri, ci).setToolTip("\n".join(e["unterschiede"]))

        lay.addWidget(tbl)

        btn_row = QHBoxLayout()
        btn_sm = QPushButton("📄  SM öffnen")
        btn_sm.setEnabled(bool(erg.pfad_staerke and os.path.isfile(erg.pfad_staerke)))
        btn_sm.clicked.connect(lambda: _oeffne_datei(erg.pfad_staerke))
        btn_sm.setStyleSheet(
            f"QPushButton{{background:#2980b9;color:white;border:none;border-radius:4px;"
            f"padding:6px 14px;font-size:11px;font-weight:bold;}}"
            f"QPushButton:hover{{background:#1a6ea8;}}"
            f"QPushButton:disabled{{background:#ccc;color:#888;}}"
        )
        btn_dp = QPushButton("📋  Dienstplan öffnen")
        btn_dp.setEnabled(bool(erg.pfad_dienstplan and os.path.isfile(erg.pfad_dienstplan)))
        btn_dp.clicked.connect(lambda: _oeffne_datei(erg.pfad_dienstplan))
        btn_dp.setStyleSheet(
            f"QPushButton{{background:#27ae60;color:white;border:none;border-radius:4px;"
            f"padding:6px 14px;font-size:11px;font-weight:bold;}}"
            f"QPushButton:hover{{background:#1e8449;}}"
            f"QPushButton:disabled{{background:#ccc;color:#888;}}"
        )
        btn_close = QPushButton("Schließen")
        btn_close.clicked.connect(self.accept)
        btn_close.setStyleSheet(
            f"QPushButton{{background:{FIORI_BLUE};color:white;border:none;border-radius:4px;"
            f"padding:6px 18px;font-size:12px;font-weight:bold;}}"
            f"QPushButton:hover{{background:#1a5276;}}"
        )
        btn_row.addWidget(btn_sm)
        btn_row.addWidget(btn_dp)
        btn_row.addStretch()
        btn_row.addWidget(btn_close)
        lay.addLayout(btn_row)

    def _kontext_menu(self, tbl: QTableWidget, pos):
        menu = QMenu(self)
        a_sm = menu.addAction("📄  Stärkemeldung öffnen")
        a_sm.setEnabled(bool(self._erg.pfad_staerke and os.path.isfile(self._erg.pfad_staerke)))
        a_dp = menu.addAction("📋  Tagesdienstplan öffnen")
        a_dp.setEnabled(bool(self._erg.pfad_dienstplan and os.path.isfile(self._erg.pfad_dienstplan)))
        action = menu.exec(tbl.viewport().mapToGlobal(pos))
        if action == a_sm:
            _oeffne_datei(self._erg.pfad_staerke)
        elif action == a_dp:
            _oeffne_datei(self._erg.pfad_dienstplan)


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _oeffne_datei(pfad: str):
    """Öffnet eine Datei mit dem systemseitigen Standardprogramm."""
    if pfad and os.path.isfile(pfad):
        QDesktopServices.openUrl(QUrl.fromLocalFile(pfad))


# ── Haupt-Widget ───────────────────────────────────────────────────────────────

class WorkflowWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._monat:            str       = ""
        self._staerke_pfade:    list[str] = []
        self._dienstplan_pfade: list[str] = []
        self._ergebnisse:       list[AbgleichErgebnis] = []
        self._thread: _LadeThread | None = None
        # Refs auf Dateilisten-Layouts (gesetzt in _build_datei_gruppe)
        self._sm_vlay:     QVBoxLayout | None = None
        self._sm_lbl_leer: QLabel      | None = None
        self._dp_vlay:     QVBoxLayout | None = None
        self._dp_lbl_leer: QLabel      | None = None
        self._setup_ui()
        # Gespeicherte Monate laden und zuletzt genutzten vorauswählen
        self._lade_monate_combo(auto_select=True)

    # ── UI ─────────────────────────────────────────────────────────────────────

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(16, 10, 16, 10)
        root.setSpacing(10)

        # Titel
        title = QLabel("⚙️  Workflow – Stärkemeldung ↔ Dienstplan Abgleich")
        title.setFont(QFont("Segoe UI", 18, QFont.Weight.Light))
        title.setStyleSheet(f"color: {FIORI_TEXT};")
        root.addWidget(title)

        hint = QLabel(
            "Lade Stärkemeldungen (Word) und Tagesdienstpläne (Excel) um sie automatisch abzugleichen. "
            "Es wird geprüft ob Namen, Dienste und Zeiten übereinstimmen."
        )
        hint.setStyleSheet("color: #666; font-size: 11px;")
        hint.setWordWrap(True)
        root.addWidget(hint)

        # ── Monat-Auswahl ──────────────────────────────────────────────────────
        monat_frame = QFrame()
        monat_frame.setStyleSheet(
            "QFrame{background:#eaf4fb;border:1px solid #aed6f1;border-radius:6px;padding:2px;}"
        )
        monat_lay = QHBoxLayout(monat_frame)
        monat_lay.setContentsMargins(10, 6, 10, 6)
        monat_lay.setSpacing(10)

        monat_lbl = QLabel("📅  Monat:")
        monat_lbl.setStyleSheet("font-size: 12px; font-weight: bold; color: #1a5276; background: transparent; border: none;")
        monat_lay.addWidget(monat_lbl)

        self._monat_combo = QComboBox()
        self._monat_combo.setMinimumWidth(200)
        self._monat_combo.setStyleSheet(
            "QComboBox{font-size:12px;padding:3px 8px;border:1px solid #aed6f1;"
            "border-radius:4px;background:white;}"
        )
        self._monat_combo.currentIndexChanged.connect(self._monat_gewaehlt)
        monat_lay.addWidget(self._monat_combo)

        btn_neu = QPushButton("📅  Neuer Monat")
        btn_neu.setMinimumHeight(32)
        btn_neu.setStyleSheet(self._btn_style("#27ae60", "#1e8449"))
        btn_neu.clicked.connect(self._neuer_monat)
        monat_lay.addWidget(btn_neu)

        self._btn_monat_loeschen = QPushButton("🗑  Monat entfernen")
        self._btn_monat_loeschen.setMinimumHeight(32)
        self._btn_monat_loeschen.setEnabled(False)
        self._btn_monat_loeschen.setStyleSheet(self._btn_style("#e74c3c", "#c0392b"))
        self._btn_monat_loeschen.clicked.connect(self._monat_loeschen)
        monat_lay.addWidget(self._btn_monat_loeschen)

        self._monat_info_lbl = QLabel("")
        self._monat_info_lbl.setStyleSheet("color:#1a5276;font-size:11px;background:transparent;border:none;")
        monat_lay.addWidget(self._monat_info_lbl)
        monat_lay.addStretch()
        root.addWidget(monat_frame)

        # Dateiauswahl-Bereich
        datei_row = QHBoxLayout()
        datei_row.setSpacing(12)

        datei_row.addWidget(self._build_datei_gruppe(
            "📄  Stärkemeldungen (Word .docx)",
            "staerke",
            _DEFAULT_STAERKE,
            ["Word-Dokumente (*.docx *.doc)"],
        ))
        datei_row.addWidget(self._build_datei_gruppe(
            "📋  Tagesdienstpläne (Excel .xlsx)",
            "dienstplan",
            _DEFAULT_DIENSTPLAN,
            ["Excel-Tabellen (*.xlsx *.xls)"],
        ))
        root.addLayout(datei_row)

        # Aktions-Zeile
        btn_row = QHBoxLayout()
        self._btn_start = QPushButton("🔄  Abgleich starten")
        self._btn_start.setMinimumHeight(40)
        self._btn_start.setEnabled(False)
        self._btn_start.setStyleSheet(self._btn_style(FIORI_BLUE, "#1a5276"))
        self._btn_start.clicked.connect(self._starte_abgleich)
        btn_row.addWidget(self._btn_start)

        self._btn_reset = QPushButton("🗑  Zurücksetzen")
        self._btn_reset.setMinimumHeight(40)
        self._btn_reset.setStyleSheet(self._btn_style("#7f8c8d", "#5d6d7e"))
        self._btn_reset.clicked.connect(self._reset)
        btn_row.addWidget(self._btn_reset)
        btn_row.addStretch()
        root.addLayout(btn_row)

        # Fortschrittsbalken
        self._progress = QProgressBar()
        self._progress.setVisible(False)
        self._progress.setRange(0, 0)
        self._progress.setFixedHeight(6)
        self._progress.setStyleSheet(
            f"QProgressBar{{border:none;border-radius:3px;background:#e0e0e0;}}"
            f"QProgressBar::chunk{{background:{FIORI_BLUE};border-radius:3px;}}"
        )
        root.addWidget(self._progress)

        self._status_lbl = QLabel("")
        self._status_lbl.setStyleSheet("color: #555; font-size: 11px;")
        root.addWidget(self._status_lbl)

        # Ergebnis-Tabelle
        grp_erg = QGroupBox("📊  Abgleich-Ergebnisse")
        grp_erg.setStyleSheet(
            f"QGroupBox{{border:1px solid #c8d0d8;border-radius:6px;margin-top:14px;"
            f"padding-top:6px;font-size:11px;font-weight:bold;color:{FIORI_TEXT};}}"
            f"QGroupBox::title{{subcontrol-origin:margin;subcontrol-position:top left;"
            f"left:10px;padding:0 6px;color:{FIORI_BLUE};}}"
        )
        erg_lay = QVBoxLayout(grp_erg)

        self._erg_table = QTableWidget()
        self._erg_table.setColumnCount(8)
        self._erg_table.setHorizontalHeaderLabels([
            "Stärkemeldung", "Dienstplan",
            "✅ Identisch", "⚠️ Abweichung",
            "📄 Nur Stärke", "📋 Nur Dienstplan",
            "📝 Carmen", "💬 Notiz",
        ])
        self._erg_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self._erg_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        for ci in range(2, 6):
            self._erg_table.horizontalHeader().setSectionResizeMode(
                ci, QHeaderView.ResizeMode.ResizeToContents
            )
        self._erg_table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)
        self._erg_table.horizontalHeader().setSectionResizeMode(7, QHeaderView.ResizeMode.Stretch)
        self._erg_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._erg_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._erg_table.setAlternatingRowColors(True)
        self._erg_table.verticalHeader().setVisible(False)
        self._erg_table.setStyleSheet("font-size: 11px;")
        self._erg_table.doubleClicked.connect(self._detail_zeigen)
        self._erg_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._erg_table.customContextMenuRequested.connect(self._kontext_menu_erg)

        # Warn-Banner für fehlende/überzählige Dateien (anfangs versteckt)
        self._warn_frame = QFrame()
        self._warn_frame.setVisible(False)
        self._warn_frame.setStyleSheet(
            "QFrame{background:#fff3cd;border:1px solid #ffc107;"
            "border-radius:6px;padding:2px;}"
        )
        warn_lay = QVBoxLayout(self._warn_frame)
        warn_lay.setContentsMargins(10, 6, 10, 6)
        warn_lay.setSpacing(2)
        self._warn_lbl = QLabel()
        self._warn_lbl.setStyleSheet("color:#856404;font-size:11px;")
        self._warn_lbl.setWordWrap(True)
        warn_lay.addWidget(self._warn_lbl)
        erg_lay.addWidget(self._warn_frame)

        erg_lay.addWidget(self._erg_table)

        hint_detail = QLabel("💡  Doppelklick für Details  |  Rechtsklick für Menü (Dateien öffnen, Carmen, Notiz)")
        hint_detail.setStyleSheet("color: #888; font-size: 10px;")
        erg_lay.addWidget(hint_detail)

        root.addWidget(grp_erg, stretch=1)

    def _build_datei_gruppe(
        self,
        titel: str,
        key: str,
        default_dir: str,
        filter_list: list[str],
    ) -> QGroupBox:
        grp = QGroupBox(titel)
        grp.setStyleSheet(
            f"QGroupBox{{border:1px solid #c8d0d8;border-radius:6px;margin-top:14px;"
            f"padding-top:6px;font-size:11px;font-weight:bold;color:{FIORI_TEXT};}}"
            f"QGroupBox::title{{subcontrol-origin:margin;subcontrol-position:top left;"
            f"left:10px;padding:0 6px;color:{FIORI_BLUE};}}"
        )
        lay = QVBoxLayout(grp)
        lay.setSpacing(4)

        # Dateiliste
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setFixedHeight(110)
        scroll.setStyleSheet("background: #fafafa; border: 1px solid #e0e0e0; border-radius: 4px;")
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        vlay = QVBoxLayout(container)
        vlay.setContentsMargins(4, 4, 4, 4)
        vlay.setSpacing(2)
        scroll.setWidget(container)
        lay.addWidget(scroll)

        lbl_leer = QLabel("Keine Dateien ausgewählt")
        lbl_leer.setStyleSheet("color: #aaa; font-size: 11px; font-style: italic;")
        vlay.addWidget(lbl_leer)
        vlay.addStretch()

        # Refs für späteres Aktualisieren durch Monat-Selektor
        if key == "staerke":
            self._sm_vlay     = vlay
            self._sm_lbl_leer = lbl_leer
        else:
            self._dp_vlay     = vlay
            self._dp_lbl_leer = lbl_leer

        # Buttons
        btn_lay = QHBoxLayout()
        btn_laden = QPushButton("📂  Dateien laden")
        btn_laden.setMinimumHeight(32)
        btn_laden.setStyleSheet(self._btn_style("#2980b9", "#1a6ea8"))
        btn_laden.clicked.connect(lambda: self._lade_dateien(key, default_dir, filter_list, vlay, lbl_leer))
        btn_lay.addWidget(btn_laden)

        btn_clear = QPushButton("✕")
        btn_clear.setFixedSize(32, 32)
        btn_clear.setToolTip("Auswahl leeren")
        btn_clear.setStyleSheet(self._btn_style("#e74c3c", "#c0392b"))
        btn_clear.clicked.connect(lambda: self._clear_dateien(key, vlay, lbl_leer))
        btn_lay.addWidget(btn_clear)
        lay.addLayout(btn_lay)

        return grp

    # ── Dateiauswahl ────────────────────────────────────────────────────────────

    def _lade_dateien(
        self,
        key: str,
        default_dir: str,
        filter_list: list[str],
        vlay: QVBoxLayout,
        lbl_leer: QLabel,
    ):
        # Standard-Dir anlegen falls nicht vorhanden
        import os
        start = default_dir if os.path.isdir(default_dir) else str(Path.home())

        dateien, _ = QFileDialog.getOpenFileNames(
            self,
            "Dateien auswählen",
            start,
            ";;".join(filter_list) + ";;Alle Dateien (*)",
        )
        if not dateien:
            return

        if key == "staerke":
            self._staerke_pfade = dateien
        else:
            self._dienstplan_pfade = dateien

        self._aktualisiere_dateiliste(vlay, lbl_leer, dateien)
        self._update_start_btn()
        self._speichere_session_wenn_monat()

    def _clear_dateien(self, key: str, vlay: QVBoxLayout, lbl_leer: QLabel):
        if key == "staerke":
            self._staerke_pfade = []
        else:
            self._dienstplan_pfade = []
        self._aktualisiere_dateiliste(vlay, lbl_leer, [])
        self._update_start_btn()
        self._speichere_session_wenn_monat()

    def _speichere_session_wenn_monat(self) -> None:
        """Schreibt die aktuellen Dateilisten in die DB, wenn ein Monat gewählt ist."""
        if self._monat:
            speichere_session(self._monat, self._staerke_pfade, self._dienstplan_pfade)
            self._monat_info_lbl.setText(
                f"💾 Gespeichert – {len(self._staerke_pfade)} SM | {len(self._dienstplan_pfade)} DP"
            )

    def _aktualisiere_dateiliste(self, vlay: QVBoxLayout, lbl_leer: QLabel, dateien: list[str]):
        # Alle alten Widgets entfernen
        while vlay.count():
            item = vlay.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not dateien:
            lbl_leer2 = QLabel("Keine Dateien ausgewählt")
            lbl_leer2.setStyleSheet("color: #aaa; font-size: 11px; font-style: italic;")
            vlay.addWidget(lbl_leer2)
        else:
            for pfad in dateien:
                lbl = QLabel(f"📄  {Path(pfad).name}")
                lbl.setStyleSheet("color: #2c3e50; font-size: 11px;")
                lbl.setToolTip(pfad)
                vlay.addWidget(lbl)
        vlay.addStretch()

    def _update_start_btn(self):
        self._btn_start.setEnabled(
            bool(self._staerke_pfade) and bool(self._dienstplan_pfade)
        )

    # ── Abgleich ────────────────────────────────────────────────────────────────

    def _starte_abgleich(self):
        if self._thread and self._thread.isRunning():
            return

        self._erg_table.setRowCount(0)
        self._progress.setVisible(True)
        self._status_lbl.setText("Wird geladen …")
        self._btn_start.setEnabled(False)

        self._thread = _LadeThread(self._staerke_pfade, self._dienstplan_pfade)
        self._thread.fortschritt.connect(self._status_lbl.setText)
        self._thread.fertig.connect(self._zeige_ergebnisse)
        self._thread.fehler.connect(self._zeige_fehler)
        self._thread.start()

    def _zeige_ergebnisse(self, ergebnisse: list):
        self._ergebnisse = ergebnisse
        self._progress.setVisible(False)
        self._btn_start.setEnabled(True)
        self._status_lbl.setText(
            f"Abgleich abgeschlossen – {len(ergebnisse)} Paar(e) verglichen."
        )

        # Warnungen: SM ohne DP / DP ohne SM
        fehlende_dp = [
            e.datei_staerke
            for e in ergebnisse
            if e.datei_dienstplan.startswith("(kein DP")
        ]
        fehlende_sm = [
            e.datei_dienstplan
            for e in ergebnisse
            if e.datei_staerke in ("–", "")
        ]
        warn_zeilen = []
        if fehlende_dp:
            warn_zeilen.append(
                f"⚠️  Kein Dienstplan gefunden für: {', '.join(fehlende_dp)}"
            )
        if fehlende_sm:
            warn_zeilen.append(
                f"⚠️  Keine Stärkemeldung gefunden für: {', '.join(fehlende_sm)}"
            )
        if warn_zeilen:
            self._warn_lbl.setText("\n".join(warn_zeilen))
            self._warn_frame.setVisible(True)
        else:
            self._warn_frame.setVisible(False)

        self._erg_table.setRowCount(len(ergebnisse))
        for ri, erg in enumerate(ergebnisse):
            hat_fehler = erg.hat_fehler
            bg = QColor("#fff3cd") if hat_fehler else QColor("#f0fff0")

            # DB-Eintrag laden (Carmen + Notiz)
            db_entry = lade_eintrag(erg.datum, erg.datei_staerke)
            ist_carmen = bool(db_entry.get("abgeglichen_carmen"))
            notiz_text = db_entry.get("notiz") or ""

            vals = [
                erg.datei_staerke,
                erg.datei_dienstplan,
                str(len(erg.ok)),
                str(len(erg.abweichung)),
                str(len(erg.nur_staerke)),
                str(len(erg.nur_dienstplan)),
            ]
            for ci, v in enumerate(vals):
                item = QTableWidgetItem(v)
                item.setBackground(bg)
                item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                if ci in (0, 1):
                    item.setTextAlignment(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft)
                if ci == 3 and erg.abweichung:
                    item.setForeground(QColor("#856404"))
                    item.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
                if ci == 4 and erg.nur_staerke:
                    item.setForeground(QColor("#c0392b"))
                if ci == 5 and erg.nur_dienstplan:
                    item.setForeground(QColor("#1a6ea8"))
                self._erg_table.setItem(ri, ci, item)

            # Spalte 6: Carmen – grüner Hintergrund + ✓ wenn abgeglichen, sonst gedimmtes –
            c_item = QTableWidgetItem()
            c_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if ist_carmen:
                c_item.setText("✓")
                c_item.setBackground(QColor("#d5f5e3"))
                c_item.setForeground(QColor("#1e8449"))
                c_item.setFont(QFont("Segoe UI", 14, QFont.Weight.Bold))
                ts = db_entry.get("abgeglichen_carmen_am") or ""
                c_item.setToolTip(f"Mit Carmen abgeglichen am {ts}" if ts else "Mit Carmen abgeglichen")
            else:
                c_item.setText("—")
                c_item.setBackground(QColor("#f2f3f4"))
                c_item.setForeground(QColor("#bdc3c7"))
                c_item.setFont(QFont("Segoe UI", 11))
                c_item.setToolTip("Noch nicht mit Carmen abgeglichen")
            self._erg_table.setItem(ri, 6, c_item)

            # Spalte 7: Notiz – 📝 Icon wenn vorhanden, sonst gedimmtes –
            n_item = QTableWidgetItem()
            n_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if notiz_text:
                n_item.setText("📝")
                n_item.setBackground(QColor("#fef9e7"))
                n_item.setForeground(QColor("#7d6608"))
                n_item.setFont(QFont("Segoe UI", 13))
                vorschau = notiz_text[:120] + ("…" if len(notiz_text) > 120 else "")
                n_item.setToolTip(vorschau)
            else:
                n_item.setText("—")
                n_item.setBackground(QColor("#f2f3f4"))
                n_item.setForeground(QColor("#bdc3c7"))
                n_item.setFont(QFont("Segoe UI", 11))
                n_item.setToolTip("Keine Notiz")
            self._erg_table.setItem(ri, 7, n_item)

        if ergebnisse:
            gesamt_ok   = sum(len(e.ok)          for e in ergebnisse)
            gesamt_abw  = sum(len(e.abweichung)  for e in ergebnisse)
            gesamt_ns   = sum(len(e.nur_staerke) for e in ergebnisse)
            gesamt_nd   = sum(len(e.nur_dienstplan) for e in ergebnisse)
            self._status_lbl.setText(
                f"Fertig – Gesamt: ✅ {gesamt_ok} identisch  "
                f"⚠️ {gesamt_abw} Abweichung(en)  "
                f"📄 {gesamt_ns} nur Stärke  "
                f"📋 {gesamt_nd} nur Dienstplan"
            )

    def _zeige_fehler(self, msg: str):
        self._progress.setVisible(False)
        self._btn_start.setEnabled(True)
        self._status_lbl.setText("Fehler aufgetreten.")
        QMessageBox.critical(self, "Fehler beim Abgleich", msg)

    def _detail_zeigen(self, index):
        ri = index.row()
        if 0 <= ri < len(self._ergebnisse):
            dlg = _DetailDialog(self._ergebnisse[ri], self)
            dlg.exec()

    def _kontext_menu_erg(self, pos):
        ri = self._erg_table.rowAt(pos.y())
        if ri < 0 or ri >= len(self._ergebnisse):
            return
        erg = self._ergebnisse[ri]
        db_entry = lade_eintrag(erg.datum, erg.datei_staerke)
        ist_carmen = bool(db_entry.get("abgeglichen_carmen"))

        menu = QMenu(self)
        a_detail = menu.addAction("🔍  Details anzeigen")
        menu.addSeparator()
        a_sm = menu.addAction("📄  Stärkemeldung öffnen")
        a_sm.setEnabled(bool(erg.pfad_staerke and os.path.isfile(erg.pfad_staerke)))
        a_dp = menu.addAction("📋  Tagesdienstplan öffnen")
        a_dp.setEnabled(bool(erg.pfad_dienstplan and os.path.isfile(erg.pfad_dienstplan)))
        a_beide = menu.addAction("📂  Beide Dateien öffnen")
        a_beide.setEnabled(
            bool(erg.pfad_staerke and os.path.isfile(erg.pfad_staerke))
            and bool(erg.pfad_dienstplan and os.path.isfile(erg.pfad_dienstplan))
        )
        menu.addSeparator()
        a_carmen = menu.addAction(
            "✖  Carmen-Abgleich aufheben" if ist_carmen else "✓  Mit Carmen abgeglichen markieren"
        )
        a_notiz = menu.addAction("💬  Notiz bearbeiten")

        action = menu.exec(self._erg_table.viewport().mapToGlobal(pos))
        if action == a_detail:
            dlg = _DetailDialog(erg, self)
            dlg.exec()
        elif action == a_sm:
            _oeffne_datei(erg.pfad_staerke)
        elif action == a_dp:
            _oeffne_datei(erg.pfad_dienstplan)
        elif action == a_beide:
            _oeffne_datei(erg.pfad_staerke)
            _oeffne_datei(erg.pfad_dienstplan)
        elif action == a_carmen:
            speichere_eintrag(
                datum=erg.datum,
                sm_datei=erg.datei_staerke,
                dp_datei=erg.datei_dienstplan,
                abgeglichen_carmen=not ist_carmen,
            )
            self._aktualisiere_zeile(ri, erg)
        elif action == a_notiz:
            self._notiz_bearbeiten(ri, erg)

    def _aktualisiere_zeile(self, ri: int, erg) -> None:
        """Liest DB-Daten neu und aktualisiert Spalten 6+7 der angegebenen Zeile."""
        db_entry = lade_eintrag(erg.datum, erg.datei_staerke)
        ist_carmen = bool(db_entry.get("abgeglichen_carmen"))
        notiz_text = db_entry.get("notiz") or ""

        item6 = self._erg_table.item(ri, 6) or QTableWidgetItem()
        item6.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        if ist_carmen:
            item6.setText("✓")
            item6.setBackground(QColor("#d5f5e3"))
            item6.setForeground(QColor("#1e8449"))
            item6.setFont(QFont("Segoe UI", 14, QFont.Weight.Bold))
            ts = db_entry.get("abgeglichen_carmen_am") or ""
            item6.setToolTip(f"Mit Carmen abgeglichen am {ts}" if ts else "Mit Carmen abgeglichen")
        else:
            item6.setText("—")
            item6.setBackground(QColor("#f2f3f4"))
            item6.setForeground(QColor("#bdc3c7"))
            item6.setFont(QFont("Segoe UI", 11))
            item6.setToolTip("Noch nicht mit Carmen abgeglichen")
        self._erg_table.setItem(ri, 6, item6)

        item7 = self._erg_table.item(ri, 7) or QTableWidgetItem()
        item7.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        if notiz_text:
            item7.setText("📝")
            item7.setBackground(QColor("#fef9e7"))
            item7.setForeground(QColor("#7d6608"))
            item7.setFont(QFont("Segoe UI", 13))
            vorschau = notiz_text[:120] + ("…" if len(notiz_text) > 120 else "")
            item7.setToolTip(vorschau)
        else:
            item7.setText("—")
            item7.setBackground(QColor("#f2f3f4"))
            item7.setForeground(QColor("#bdc3c7"))
            item7.setFont(QFont("Segoe UI", 11))
            item7.setToolTip("Keine Notiz")
        self._erg_table.setItem(ri, 7, item7)

    def _notiz_bearbeiten(self, ri: int, erg) -> None:
        """Öffnet einen Dialog zum Bearbeiten der Notiz für diesen Tag."""
        db_entry = lade_eintrag(erg.datum, erg.datei_staerke)
        alte_notiz = db_entry.get("notiz") or ""

        dlg = QDialog(self)
        dlg.setWindowTitle(f"📝  Notiz – {erg.datei_staerke}")
        dlg.resize(480, 260)
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(14, 12, 14, 12)

        lbl = QLabel(f"Notiz für <b>{erg.datei_staerke}</b>:")
        lbl.setStyleSheet("font-size: 11px;")
        lay.addWidget(lbl)

        edit = QTextEdit()
        edit.setPlainText(alte_notiz)
        edit.setPlaceholderText("Hier Notiz eingeben …")
        edit.setStyleSheet("font-size: 11px; border: 1px solid #c8d0d8; border-radius: 4px; padding: 4px;")
        lay.addWidget(edit)

        btn_row = QHBoxLayout()
        btn_ok = QPushButton("💾  Speichern")
        btn_ok.setMinimumHeight(34)
        btn_ok.setStyleSheet(
            f"QPushButton{{background:{FIORI_BLUE};color:white;border:none;border-radius:4px;"
            f"padding:5px 14px;font-size:11px;font-weight:bold;}}"
            f"QPushButton:hover{{background:#1a5276;}}"
        )
        btn_abbruch = QPushButton("Abbrechen")
        btn_abbruch.setMinimumHeight(34)
        btn_abbruch.setStyleSheet(
            "QPushButton{background:#7f8c8d;color:white;border:none;border-radius:4px;"
            "padding:5px 14px;font-size:11px;font-weight:bold;}"
            "QPushButton:hover{background:#5d6d7e;}"
        )
        btn_loeschen = QPushButton("🗑  Notiz löschen")
        btn_loeschen.setMinimumHeight(34)
        btn_loeschen.setStyleSheet(
            "QPushButton{background:#e74c3c;color:white;border:none;border-radius:4px;"
            "padding:5px 14px;font-size:11px;font-weight:bold;}"
            "QPushButton:hover{background:#c0392b;}"
        )
        btn_row.addWidget(btn_loeschen)
        btn_row.addStretch()
        btn_row.addWidget(btn_abbruch)
        btn_row.addWidget(btn_ok)
        lay.addLayout(btn_row)

        btn_abbruch.clicked.connect(dlg.reject)
        btn_ok.clicked.connect(dlg.accept)
        btn_loeschen.clicked.connect(lambda: (edit.setPlainText(""), dlg.accept()))

        if dlg.exec() == QDialog.DialogCode.Accepted:
            neue_notiz = edit.toPlainText().strip()
            speichere_eintrag(
                datum=erg.datum,
                sm_datei=erg.datei_staerke,
                dp_datei=erg.datei_dienstplan,
                notiz=neue_notiz,
            )
            self._aktualisiere_zeile(ri, erg)

    # ── Monat-Verwaltung ────────────────────────────────────────────────────────

    _MONATSNAMEN = [
        "Januar", "Februar", "März", "April", "Mai", "Juni",
        "Juli", "August", "September", "Oktober", "November", "Dezember",
    ]

    def _monat_label(self, monat: str) -> str:
        """Wandelt 'YYYY-MM' in 'Januar 2026' um."""
        try:
            j, m = monat.split("-")
            return f"{self._MONATSNAMEN[int(m) - 1]} {j}"
        except Exception:
            return monat

    def _lade_monate_combo(self, auto_select: bool = False) -> None:
        """Füllt die ComboBox mit gespeicherten Monaten."""
        self._monat_combo.blockSignals(True)
        aktueller = self._monat  # merken
        self._monat_combo.clear()
        self._monat_combo.addItem("── Monat wählen ──", "")
        monate = alle_monate()
        for m in monate:
            self._monat_combo.addItem(self._monat_label(m), m)
        self._monat_combo.blockSignals(False)

        if auto_select and monate:
            # letzten Monat vorauswählen (neueste zuerst → Index 1)
            self._monat_combo.setCurrentIndex(1)
            self._monat_gewaehlt(1)
        elif aktueller:
            idx = self._monat_combo.findData(aktueller)
            if idx >= 0:
                self._monat_combo.blockSignals(True)
                self._monat_combo.setCurrentIndex(idx)
                self._monat_combo.blockSignals(False)

    def _neuer_monat(self) -> None:
        """Dialog zum Anlegen eines neuen Monats."""
        dlg = QDialog(self)
        dlg.setWindowTitle("📅  Neuen Monat anlegen")
        dlg.resize(300, 130)
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(16, 12, 16, 12)
        lay.setSpacing(10)

        form = QHBoxLayout()
        monat_box = QComboBox()
        for name in self._MONATSNAMEN:
            monat_box.addItem(name)
        aktuell = datetime.now()
        monat_box.setCurrentIndex(aktuell.month - 1)
        monat_box.setMinimumWidth(130)
        form.addWidget(monat_box)

        jahr_spin = QSpinBox()
        jahr_spin.setRange(2020, 2040)
        jahr_spin.setValue(aktuell.year)
        form.addWidget(jahr_spin)
        lay.addLayout(form)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        lay.addWidget(buttons)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        monat_str = f"{jahr_spin.value():04d}-{monat_box.currentIndex() + 1:02d}"
        # Session anlegen falls noch nicht vorhanden
        speichere_session(monat_str, [], [])
        self._lade_monate_combo()
        idx = self._monat_combo.findData(monat_str)
        if idx >= 0:
            self._monat_combo.setCurrentIndex(idx)
            # _monat_gewaehlt wird über Signal ausgelöst

    def _monat_gewaehlt(self, idx: int) -> None:
        """Wird aufgerufen wenn ein Monat in der ComboBox gewählt wird."""
        monat = self._monat_combo.itemData(idx) or ""
        self._monat = monat
        self._btn_monat_loeschen.setEnabled(bool(monat))

        if not monat:
            self._monat_info_lbl.setText("")
            return

        # Session aus DB laden
        import json as _json
        session = lade_session(monat)
        sm_pfade = _json.loads(session.get("sm_pfade") or "[]")
        dp_pfade = _json.loads(session.get("dp_pfade") or "[]")

        # Nicht mehr existierende Pfade herausfiltern
        sm_pfade = [p for p in sm_pfade if os.path.isfile(p)]
        dp_pfade = [p for p in dp_pfade if os.path.isfile(p)]

        self._staerke_pfade    = sm_pfade
        self._dienstplan_pfade = dp_pfade

        # Dateilisten-UI aktualisieren
        if self._sm_vlay and self._sm_lbl_leer:
            self._aktualisiere_dateiliste(self._sm_vlay, self._sm_lbl_leer, sm_pfade)
        if self._dp_vlay and self._dp_lbl_leer:
            self._aktualisiere_dateiliste(self._dp_vlay, self._dp_lbl_leer, dp_pfade)
        self._update_start_btn()

        if sm_pfade and dp_pfade:
            self._monat_info_lbl.setText(
                f"✅ {len(sm_pfade)} SM | {len(dp_pfade)} DP – starte Abgleich …"
            )
            self._starte_abgleich()
        elif sm_pfade or dp_pfade:
            self._monat_info_lbl.setText(
                f"📂 {len(sm_pfade)} SM, {len(dp_pfade)} DP geladen – bitte fehlende Dateien laden"
            )
        else:
            self._monat_info_lbl.setText("Noch keine Dateien geladen")

    def _monat_loeschen(self) -> None:
        """Entfernt den aktuellen Monat aus der Session-DB (Carmen/Notizen bleiben)."""
        if not self._monat:
            return
        antwort = QMessageBox.question(
            self,
            "Monat entfernen",
            f"Soll der Monat <b>{self._monat_label(self._monat)}</b> aus der Liste entfernt werden?\n\n"
            "Carmen-Abgleich-Status und Notizen bleiben erhalten.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if antwort != QMessageBox.StandardButton.Yes:
            return
        loesche_session(self._monat)
        self._monat = ""
        self._lade_monate_combo()
        # UI leeren
        self._staerke_pfade    = []
        self._dienstplan_pfade = []
        self._ergebnisse       = []
        self._erg_table.setRowCount(0)
        self._warn_frame.setVisible(False)
        self._status_lbl.setText("")
        self._update_start_btn()
        if self._sm_vlay and self._sm_lbl_leer:
            self._aktualisiere_dateiliste(self._sm_vlay, self._sm_lbl_leer, [])
        if self._dp_vlay and self._dp_lbl_leer:
            self._aktualisiere_dateiliste(self._dp_vlay, self._dp_lbl_leer, [])

    def _reset(self):
        self._staerke_pfade = []
        self._dienstplan_pfade = []
        self._ergebnisse = []
        self._erg_table.setRowCount(0)
        self._warn_frame.setVisible(False)
        self._status_lbl.setText("")
        self._monat_info_lbl.setText("Dateien zurückgesetzt – bitte neu laden")
        self._update_start_btn()
        if self._sm_vlay and self._sm_lbl_leer:
            self._aktualisiere_dateiliste(self._sm_vlay, self._sm_lbl_leer, [])
        if self._dp_vlay and self._dp_lbl_leer:
            self._aktualisiere_dateiliste(self._dp_vlay, self._dp_lbl_leer, [])

    # ── Styles ─────────────────────────────────────────────────────────────────

    def _btn_style(self, bg: str, hover: str) -> str:
        return (
            f"QPushButton{{background:{bg};color:white;border:none;border-radius:4px;"
            f"padding:5px 14px;font-size:11px;font-weight:bold;}}"
            f"QPushButton:hover{{background:{hover};}}"
            f"QPushButton:disabled{{background:#ccc;color:#888;}}"
        )

    def refresh(self):
        pass
